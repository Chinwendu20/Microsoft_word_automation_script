from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate
import os
import sys

folder = sys.argv[1]


current_working_directory=os.getcwd()

lis =['Job Title', 'Job Grade', 'Business Unit', 'Department', 'Job Type',
'Location', 'Objective/Purpose of Job', 'REPORTING RELATIONSHIPS', 'Reports to',
'Supervises', 'JOB DUTIES / RESPONSIBILITIES / ACCOUNTABILITIES', 'Internally Relates with',
'Externally Relates with', 'Budgetary Responsibility']




full_path_of_folder=os.path.join(current_working_directory, folder)

list_of_files_in_dir=os.listdir(full_path_of_folder)

try:

	os.mkdir("{}/new_JDS".format(current_working_directory))
except FileExistsError:
	pass

try:

	os.mkdir("{}/bad_files".format(current_working_directory))
except FileExistsError:
	pass



for file in list_of_files_in_dir:

	document=Document("{}/{}".format(full_path_of_folder, file))

	Job_title=document.tables[0].rows[0].cells[3].text
	

	reports_to=document.tables[0].rows[5].cells[3].text


	budgetary_responsibility=document.tables[0].rows[8].cells[3].text
	

	duties_and_responsiblities=document.tables[0].rows[10].cells[0].text.strip().split('\n')

	list_duties_and_responsiblities=[]
	for items in duties_and_responsiblities:
		if items.strip():
			list_duties_and_responsiblities.append(items)

		
		
	job_purpose=document.tables[0].rows[3].cells[3].text


	job_grade=document.tables[0].rows[0].cells[8].text


	department=document.tables[0].rows[1].cells[8].text

	business_unit=document.tables[0].rows[1].cells[3].text


	job_type=document.tables[0].rows[2].cells[3].text

	location=document.tables[0].rows[2].cells[8].text

	supervises=document.tables[0].rows[5].cells[8].text

	internally_relates_with=document.tables[0].rows[7].cells[3].text

	externally_relates_with=document.tables[0].rows[7].cells[8].text




	education_and_requirements=document.tables[0].rows[12].cells[0].text.strip().split('\n')
		
	list_education_and_requirements=[]
	for items in education_and_requirements:
		if items.strip():
			list_education_and_requirements.append(items)



	context={


	'Job_title':
	Job_title,

	'reports_to':
	reports_to,

	'list_duties_and_responsiblities':
	list_duties_and_responsiblities,

	'list_education_and_requirements':
	list_education_and_requirements,

	'job_purpose':job_purpose,

	'budgetary_responsibility':budgetary_responsibility,

	'externally_relates_with': externally_relates_with,

	'internally_relates_with': internally_relates_with,

	'supervises': supervises,

	'job_type' : job_type,

	'location' :  location,

	'business_unit' : business_unit,

	'department': department,

	'job_grade' : job_grade


	}
		
	new_doc=DocxTemplate("template.docx")
	new_doc.render(context)

	for keys in context:
		for item in lis:
			if context[keys]==item:
				new_doc.save('bad_files/new_{}'.format(file))
				break
	if 'new_{}'.format(file) not in os.listdir(os.path.join(current_working_directory, 'bad_files')):
		new_doc.save('new_JDS/new_{}'.format(file))

print('All done, you will find your properly created files in the new_JDS folder, located in your current directory')
print('='*20)
print('You would find the bad ones in bad_files folder, located in your current working directory')





