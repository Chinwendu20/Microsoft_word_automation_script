from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate
import os
import sys

folder = sys.argv[1]


current_working_directory=os.getcwd()

full_path_of_folder=os.path.join(current_working_directory, folder)

list_of_files_in_dir=os.listdir(full_path_of_folder)

try:

	os.mkdir("{}/new_JDS".format(current_working_directory))
except FileExistsError:
	pass

for file in list_of_files_in_dir:

	document=Document("{}/{}".format(full_path_of_folder, file))

	Job_title=document.tables[0].rows[0].cells[3].text
		

	reports_to=document.tables[0].rows[5].cells[3].text
	

	duties_and_responsiblities=document.tables[0].rows[10].cells[0].text.strip().split('\n')
	

	job_purpose=document.tables[0].rows[3].cells[3].text


	job_grade=document.tables[0].rows[0].cells[3].text


	department=document.tables[0].rows[1].cells[3].text
	
	business_unit=document.tables[0].rows[1].cells[1].text

	budgetary_responsibility=document.tables[0].rows[8].cells[3].text
	

	job_type=document.tables[0].rows[2].cells[1].text

	location=document.tables[0].rows[2].cells[3].text

	supervises=document.tables[0].rows[5].cells[3].text

	internally_relates_with=document.tables[0].rows[7].cells[1].text

	externally_relates_with=document.tables[0].rows[7].cells[3].text




	education_and_requirements=document.tables[0].rows[12].cells[0].text.strip().split('\n')
	

	new_doc=DocxTemplate("template.docx")
	context={


	'Job_title':
	Job_title,

	'reports_to':
	reports_to,

	'duties_and_responsiblities':
	duties_and_responsiblities,

	'education_and_requirements':
	education_and_requirements,

	'job_purpose':job_purpose,

	'budgetary_responsibility':budgetary_responsibility,

	'externally_relates_with': externally_relates_with,

	'internally_relates_with': internally_relates_with,

	'supervises': supervises

	}

	new_doc.render(context)
	new_doc.save('new_JDS/new_{}'.format(file))

print('All done, you will find your files in the new_JDS folder, located in your current directory')


job_purpose=document.tables[0].rows[3].cells[3].text


document=Document("Talent Manager x.docx")


budgetary_responsibility=document.tables[0].rows[8].cells[0].text
